import logging
import os
import customtkinter as ctk
from io import BytesIO
from typing import List, Optional

from docx import Document # From python-docx library
from docx.shared import Inches
from PIL import Image # Pillow for image handling

from src.config.config_manager import config_manager
from src.utils.file_manager import file_manager # Use centralized file manager for paths
from src.ui.dialogs import showwarning, askstring # Use centralized dialogs

logger = logging.getLogger(__name__)

def save_to_word(screenshots: List[BytesIO], titles: List[str]) -> Optional[str]:
    """
    Creates a Word document (.docx) containing the captured screenshots.

    Args:
        screenshots: A list of BytesIO objects, each containing a screenshot image.
        titles: A list of titles corresponding to each screenshot.

    Returns:
        The absolute path to the saved Word document, or None if saving failed
        or was cancelled.
    """
    if not screenshots:
        showwarning("No Screenshots", "No screenshots available to save.")
        logger.warning("save_to_word called with no screenshots.")
        return None

    # --- Get Save Location and Name ---
    doc_name = askstring("New Document", "Enter document name (without extension):")
    if not doc_name:
        logger.info("User cancelled saving Word document.")
        return None # User cancelled

    # Ensure name doesn't have extension, add .docx later
    doc_name = os.path.splitext(doc_name)[0]

    # Use FileManager to get the full save path
    try:
        # FileManager's get_save_path expects filename and extension separately
        doc_path = file_manager.get_save_path(filename=doc_name, extension="docx")
    except Exception as e:
            logger.error(f"Error getting save path from FileManager: {e}", exc_info=True)
            ctk.showerror("Path Error", f"Could not determine save path: {e}")
            return None

    # Ensure the save directory exists (FileManager __init__ should do this, but double-check)
    save_dir = os.path.dirname(doc_path)
    try:
        os.makedirs(save_dir, exist_ok=True)
    except OSError as e:
        logger.error(f"Failed to create save directory '{save_dir}': {e}")
        ctk.showerror("Directory Error", f"Could not create directory:\n{save_dir}\nError: {e}")
        return None


    # --- Create Word Document ---
    try:
        document = Document()
        document.add_heading(f"Captured Screenshots: {doc_name}", level=1)

        # Read image width from config, with fallback
        image_width_inches = config_manager.get_float("GENERAL", "image_width_inches", fallback=6.0)
        if image_width_inches <= 0:
                logger.warning(f"Invalid image_width_inches ({image_width_inches}), defaulting to 6.0.")
                image_width_inches = 6.0

        logger.info(f"Creating Word document at '{doc_path}' with {len(screenshots)} images.")

        for i, img_io in enumerate(screenshots):
            title = titles[i] if i < len(titles) else f"Screenshot {i+1}"
            logger.debug(f"Adding screenshot '{title}' to document.")

            # Add title/paragraph for the image
            document.add_paragraph(f"({i+1}) {title}", style='ListNumber') # Or use a custom style

            # Reset stream position before reading
            img_io.seek(0)

            # Add picture - python-docx needs a file-like object or path
            # Using Inches directly for width control
            try:
                    document.add_picture(img_io, width=Inches(image_width_inches))
            except Exception as img_err:
                logger.error(f"Failed to add image '{title}' to document: {img_err}", exc_info=True)
                document.add_paragraph(f"[Error adding image: {title} - {img_err}]")


            document.add_paragraph() # Add some space between images


        # --- Save Document ---
        document.save(doc_path)
        logger.info(f"Word document successfully saved: {doc_path}")
        return doc_path

    except Exception as e:
        logger.error(f"Failed to create or save Word document '{doc_path}': {e}", exc_info=True)
        ctk.showerror("Save Error", f"Could not save Word document:\n{e}")
        return None